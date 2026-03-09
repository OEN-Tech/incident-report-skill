#!/usr/bin/env python3
"""
Generate cybersecurity incident reports in the official Taiwan government format.
(個人資料侵害事故通報與紀錄表)

Produces a .docx file with:
  Part 1: Government notification form (個人資料侵害事故通報與紀錄表)
  Part 2: Detailed appendix (附錄 — 說明文件)

Usage:
    python3 generate.py --output report.docx --config config.json
    python3 generate.py --output report.docx  # uses example defaults

Requires: python-docx (`pip install python-docx`)
"""

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ---------------------------------------------------------------------------
# Example Defaults (replace with your own via config.json)
# ---------------------------------------------------------------------------

DEFAULTS = {
    "company_name": "ACME Technology Co., Ltd.",
    "receiving_agency": "數位發展部數位產業署",
    "report_date": datetime.now().strftime("%Y-%m-%d"),
    "report_time": "12:00",
    "reporter": {
        "name": "John Doe",
        "title": "Security Engineer",
        "phone": "0900-000000",
        "email": "security@example.com",
        "address": "123 Example Street, Taipei",
    },
    "incident_date": "",
    "incident_type": "其他",
    "incident_type_note": "",
    "records_affected": "",
    "general_records": 0,
    "special_records": 0,
    "cause_summary": "請參考底部附錄",
    "damage": "",
    "possible_consequences": "",
    "countermeasures": "請參考底部附錄",
    "notification_plan": "",
    "within_72_hours": True,
    "within_72_hours_reason": "",
}

DEFAULT_AUDIT_ITEMS = [
    ["GuardDuty 威脅偵測紀錄", "系統威脅偵測", "無異常"],
    ["GuardDuty IAM 異常連線偵測", "IAM 憑證異常存取", "無異常"],
    ["Security Hub 安全態勢檢查", "統一安全合規狀態", "無異常發現"],
    ["CloudTrail API 呼叫稽核", "全 API 存取紀錄", "無未授權存取"],
    ["WAF 日誌分析", "Web 應用防火牆日誌", "無異常攻擊模式"],
    ["Database 資料存取紀錄", "資料庫讀寫操作", "無異常查詢或匯出"],
    ["Function 執行日誌", "全部函式執行紀錄", "無異常調用"],
    ["Secrets Manager 存取紀錄", "機密資料存取", "無未授權存取"],
    ["KMS 金鑰使用紀錄", "加密金鑰操作與存取", "無未授權使用"],
]

DEFAULT_PROCEDURES = [
    "網路安全管理程序 — 網路分段、加密傳輸、IDS/IPS 部署規範",
    "存取控制管理程序 — 角色權限控管、最小權限原則、稽核軌跡記錄",
    "資訊安全事件管理程序 — 事件偵測、證據保全、根因分析",
    "個人資料管理程序 — 個資處理、隱私影響評估",
    "資訊安全事件通報及危機處理作業說明書 — 事件分級、通報流程、危機處理",
    "帳號及密碼管理要點 — 密碼強度、定期輪替、遠端存取雙因素驗證",
    "防火牆管理作業說明書 — Security Group 設定、預設拒絕規則、ACL 管控",
    "加密金鑰管理作業說明書 — KMS 管理、金鑰輪替、職責分離",
]

DEFAULT_SECURITY = {
    "intro": "本公司平台採用雲端原生架構，並依循業界標準建置多層次安全防護機制",
    "standards_intro": "系統之設計、開發與營運遵循以下國際標準規範：",
    "standards": [
        "PCI DSS Level 1（Payment Card Industry Data Security Standard）：支付卡產業資料安全標準",
        "ISO 27001：資訊安全管理系統（ISMS）",
        "ISO 27701：隱私資訊管理系統（PIMS）",
    ],
    "subsections": {
        "4.1 基礎架構": [
            "採用 Serverless 架構，無傳統伺服器管理風險",
            "VPC 網路隔離，多個 Security Group 分區管理",
        ],
        "4.2 加密與金鑰管理": [
            "KMS 加密金鑰管理，啟用自動輪替機制",
            "Secrets Manager 管理所有敏感憑證，避免硬編碼風險",
            "檔案儲存庫阻擋公開存取、強制 TLS 傳輸加密",
        ],
        "4.3 API 安全": [
            "強制 TLS 1.2 安全傳輸策略",
            "邊緣驗證與安全標頭注入",
            "WAF 防護，速率限制、阻擋惡意請求、Anti-DDoS",
            "Content-Security-Policy（CSP）標頭注入",
        ],
        "4.4 監控與威脅偵測": [
            "持續性威脅偵測服務",
            "統一安全態勢管理與合規檢查",
            "完整 API 呼叫稽核日誌",
            "全鏈路請求追蹤",
            "即時效能與異常監控",
        ],
        "4.5 入侵偵測與防禦": {
            "intro": "採用多層次入侵偵測與防禦架構：",
            "items": [
                "持續分析網路流量與事件日誌，自動偵測可疑活動與潛在入侵行為",
                "完整網路層可視性，可追蹤異常連線與資料傳輸行為",
                "即時攔截惡意請求，包含 SQL Injection、XSS、異常請求頻率等攻擊模式",
            ],
        },
        "4.6 存取控制": {
            "system_title": "系統層存取控制（IAM）",
            "system_items": [
                "嚴格遵循最小權限原則（Principle of Least Privilege）",
                "強制啟用 MFA 多因素驗證",
                "密碼管理：最低 12 字元、定期輪替、連續失敗鎖定帳戶",
                "定期進行帳戶盤點",
            ],
            "app_title": "應用層存取控制（RBAC）",
            "app_items": [
                "租戶隔離之角色型存取控制（Role-Based Access Control）",
                "管理後台登入強制二階段驗證（OTP）",
            ],
        },
    },
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def deep_merge(base, override):
    """Recursively merge override into base dict."""
    result = base.copy()
    for k, v in override.items():
        if k in result and isinstance(result[k], dict) and isinstance(v, dict):
            result[k] = deep_merge(result[k], v)
        else:
            result[k] = v
    return result


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top/bottom/left/right, each {sz, val, color}."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, props in kwargs.items():
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{props.get("val", "single")}" '
            f'w:sz="{props.get("sz", "4")}" w:space="0" '
            f'w:color="{props.get("color", "000000")}"/>'
        )
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_kai_text(cell, text, size=14, bold=False, alignment=None):
    """Add text in Kai (楷體) font to a table cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = "Kai"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Kai")
    return p


def bordered_cells(table):
    """Apply standard borders to all cells in a table."""
    border = {"sz": "4", "val": "single", "color": "000000"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border, bottom=border, left=border, right=border)


# ---------------------------------------------------------------------------
# Part 1: Government Form (個人資料侵害事故通報與紀錄表)
# ---------------------------------------------------------------------------

def build_form(doc, cfg):
    """Build the government notification form table."""
    rpt = cfg.get("reporter", DEFAULTS["reporter"])
    d = cfg.get("report_date", DEFAULTS["report_date"])
    parts = d.split("-")
    yr, mo, dy = (parts + ["", "", ""])[:3]
    t = cfg.get("report_time", DEFAULTS["report_time"])
    t_parts = t.split(":")
    hr, mn = (t_parts + ["12", "00"])[:2]

    inc_d = cfg.get("incident_date", "")
    inc_parts = inc_d.split("-") if inc_d else ["", "", ""]
    inc_yr, inc_mo, inc_dy = (inc_parts + ["", "", ""])[:3]

    company = cfg.get("company_name", DEFAULTS["company_name"])
    agency = cfg.get("receiving_agency", DEFAULTS["receiving_agency"])
    inc_type = cfg.get("incident_type", DEFAULTS["incident_type"])
    inc_note = cfg.get("incident_type_note", DEFAULTS["incident_type_note"])
    records = cfg.get("records_affected", DEFAULTS["records_affected"])
    gen_rec = cfg.get("general_records", 0)
    spe_rec = cfg.get("special_records", 0)

    table = doc.add_table(rows=20, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    bordered_cells(table)

    # Row 0: Title
    merged = table.cell(0, 0).merge(table.cell(0, 1)).merge(table.cell(0, 2))
    add_kai_text(merged, "個人資料侵害事故通報與紀錄表 ", size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Rows 1-6: Left = company/agency info, Right = reporter info
    left_merged = table.cell(1, 0).merge(table.cell(6, 0))
    add_kai_text(left_merged, f"事業名稱\n{company}\n通報機關\n{agency}", size=14)

    r1 = table.cell(1, 1).merge(table.cell(1, 2))
    add_kai_text(r1, f"通報時間： {yr} 年  {mo}月 {dy} 日  {hr}時 {mn} 分 ", size=14)

    r2 = table.cell(2, 1).merge(table.cell(2, 2))
    add_kai_text(r2, f"通報人：    {rpt.get('name', '')} 簽名(蓋章) ", size=14)

    r3 = table.cell(3, 1).merge(table.cell(3, 2))
    add_kai_text(r3, f"職稱：       {rpt.get('title', '')}", size=14)

    r4 = table.cell(4, 1).merge(table.cell(4, 2))
    add_kai_text(r4, f"電話：       {rpt.get('phone', '')}", size=14)

    r5 = table.cell(5, 1).merge(table.cell(5, 2))
    add_kai_text(r5, f"Email：      {rpt.get('email', '')}", size=14)

    r6 = table.cell(6, 1).merge(table.cell(6, 2))
    add_kai_text(r6, f"地址：       {rpt.get('address', '')}", size=14)

    # Row 7: Incident date
    add_kai_text(table.cell(7, 0), "事件發生時間", size=14)
    r7 = table.cell(7, 1).merge(table.cell(7, 2))
    add_kai_text(r7, f"{inc_yr} 年 {inc_mo}月 {inc_dy} 日", size=14)

    # Rows 8-13: Incident type + records affected
    type_left = table.cell(8, 0).merge(table.cell(13, 0))
    add_kai_text(type_left, "事件發生種類", size=14)

    types = ["竊取", "洩漏", "竄改", "毀損", "滅失", "其他"]
    for i, t_name in enumerate(types):
        row_idx = 8 + i
        cell = table.cell(row_idx, 1)
        check = "■" if inc_type == t_name else "□"
        text = f"{check}{t_name}"
        if t_name == "其他" and inc_type == "其他" and inc_note:
            text += f" | {inc_note}"
        add_kai_text(cell, text, size=14)

    rec_top = table.cell(8, 2).merge(table.cell(10, 2))
    add_kai_text(rec_top, f"個資侵害之總筆數(大約)  | {records}", size=14)

    rec_bot = table.cell(11, 2).merge(table.cell(13, 2))
    add_kai_text(rec_bot, f"□一般個_{gen_rec}_筆  | □特種個_{spe_rec}_筆", size=14)

    # Row 14: Cause summary
    add_kai_text(table.cell(14, 0), "發生原因及事件摘要", size=14)
    r14 = table.cell(14, 1).merge(table.cell(14, 2))
    add_kai_text(r14, cfg.get("cause_summary", DEFAULTS["cause_summary"]), size=14)

    # Row 15: Damage
    add_kai_text(table.cell(15, 0), "損害狀況", size=14)
    r15 = table.cell(15, 1).merge(table.cell(15, 2))
    add_kai_text(r15, cfg.get("damage", DEFAULTS["damage"]), size=14)

    # Row 16: Possible consequences
    add_kai_text(table.cell(16, 0), "個資侵害可能結果", size=14)
    r16 = table.cell(16, 1).merge(table.cell(16, 2))
    add_kai_text(r16, cfg.get("possible_consequences", DEFAULTS["possible_consequences"]), size=14)

    # Row 17: Countermeasures
    add_kai_text(table.cell(17, 0), "擬採取之因應措施", size=14)
    r17 = table.cell(17, 1).merge(table.cell(17, 2))
    add_kai_text(r17, cfg.get("countermeasures", DEFAULTS["countermeasures"]), size=14)

    # Row 18: Notification plan
    add_kai_text(table.cell(18, 0), "擬採通知當事人之時\n間及方式", size=14)
    r18 = table.cell(18, 1).merge(table.cell(18, 2))
    add_kai_text(r18, cfg.get("notification_plan", DEFAULTS["notification_plan"]), size=14)

    # Row 19: 72-hour notification
    add_kai_text(table.cell(19, 0), "是否於發現個資外洩\n時起算七十二小時內\n通報", size=14)
    r19 = table.cell(19, 1).merge(table.cell(19, 2))
    w72 = cfg.get("within_72_hours", True)
    reason = cfg.get("within_72_hours_reason", DEFAULTS["within_72_hours_reason"])
    yes_check = "■" if w72 else "□"
    no_check = "□" if w72 else "■"
    add_kai_text(r19, f"{yes_check}是  {no_check}否，理由：{reason}", size=14)


# ---------------------------------------------------------------------------
# Part 2: Appendix (附錄)
# ---------------------------------------------------------------------------

def build_appendix(doc, cfg):
    """Build the detailed explanation appendix."""
    app = cfg.get("appendix", {})
    company = cfg.get("company_name", DEFAULTS["company_name"])
    d = cfg.get("report_date", DEFAULTS["report_date"])
    parts = d.split("-")
    yr, mo, dy = (parts + ["", "", ""])[:3]

    # Title
    title = app.get("title", "資安事件說明文件")
    h = doc.add_heading(f"附錄. {title}", level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta
    doc.add_paragraph(f"文件日期： {yr} 年 {mo} 月 {dy} 日")
    doc.add_paragraph(f"文件性質： {app.get('doc_nature', '資安事件通報說明')}")
    doc.add_paragraph(f"編製單位： {company}")

    sections = app.get("sections", {})

    # 一、事件摘要
    doc.add_heading("一、事件摘要", level=2)
    if sections.get("summary"):
        doc.add_paragraph(sections["summary"])

    # 二、與本公司之關聯
    doc.add_heading("二、與本公司之關聯", level=2)
    if sections.get("relation"):
        doc.add_paragraph(sections["relation"])
    if sections.get("relation_conclusion"):
        doc.add_paragraph(sections["relation_conclusion"])
    if sections.get("relation_details"):
        doc.add_paragraph("具體說明如下：")
        for detail in sections["relation_details"]:
            doc.add_paragraph(detail)

    # 三、事件時間軸
    doc.add_heading("三、事件時間軸", level=2)
    timeline = sections.get("timeline", [])
    if timeline:
        tl = doc.add_table(rows=1 + len(timeline), cols=2)
        tl.alignment = WD_TABLE_ALIGNMENT.CENTER
        bordered_cells(tl)
        add_kai_text(tl.cell(0, 0), "時間", size=11, bold=True)
        add_kai_text(tl.cell(0, 1), "事件", size=11, bold=True)
        for i, (time_val, event_val) in enumerate(timeline):
            add_kai_text(tl.cell(i + 1, 0), time_val, size=11)
            add_kai_text(tl.cell(i + 1, 1), event_val, size=11)

    # 四、系統安全架構說明
    doc.add_heading("四、本公司系統安全架構說明", level=2)
    sec = sections.get("security_architecture", DEFAULT_SECURITY)
    doc.add_paragraph(sec.get("intro", DEFAULT_SECURITY["intro"]))
    doc.add_paragraph(sec.get("standards_intro", DEFAULT_SECURITY["standards_intro"]))
    for std in sec.get("standards", DEFAULT_SECURITY["standards"]):
        doc.add_paragraph(std)

    subsections = sec.get("subsections", DEFAULT_SECURITY["subsections"])
    for sub_title, sub_content in subsections.items():
        if isinstance(sub_content, list):
            doc.add_heading(sub_title, level=3)
            for item in sub_content:
                doc.add_heading(item, level=3)
        elif isinstance(sub_content, dict):
            doc.add_heading(sub_title, level=3)
            if "intro" in sub_content:
                doc.add_heading(sub_content["intro"], level=3)
                for item in sub_content.get("items", []):
                    doc.add_paragraph(item)
            if "system_title" in sub_content:
                doc.add_paragraph(sub_content["system_title"])
                for item in sub_content.get("system_items", []):
                    doc.add_paragraph(item)
                doc.add_paragraph(sub_content.get("app_title", ""))
                for item in sub_content.get("app_items", []):
                    doc.add_paragraph(item)

    # 五、系統排查報告
    doc.add_heading("五、系統排查報告", level=2)
    procedures_intro = sections.get("audit_procedures_intro",
        "本公司依循資訊安全與隱私管理規範，於事件發生後立即啟動系統性排查作業，"
        "參考以下內部管理程序：")
    doc.add_paragraph(procedures_intro)
    for proc in sections.get("audit_procedures", DEFAULT_PROCEDURES):
        doc.add_paragraph(proc)

    doc.add_heading("排查結果", level=3)
    audit_items = sections.get("audit_results", DEFAULT_AUDIT_ITEMS)
    if audit_items:
        at = doc.add_table(rows=1 + len(audit_items), cols=4)
        at.alignment = WD_TABLE_ALIGNMENT.CENTER
        bordered_cells(at)
        for ci, h_text in enumerate(["項次", "排查項目", "排查範圍", "結果"]):
            add_kai_text(at.cell(0, ci), h_text, size=11, bold=True)
        for i, row_data in enumerate(audit_items):
            add_kai_text(at.cell(i + 1, 0), str(i + 1), size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            for j, val in enumerate(row_data):
                add_kai_text(at.cell(i + 1, j + 1), val, size=11)

    conclusion = sections.get("audit_conclusion",
        "排查結論：本公司系統於事件期間未發現任何異常存取、未授權操作或入侵跡象。")
    doc.add_paragraph(conclusion)

    # 六、結論
    doc.add_heading("六、結論", level=2)
    for c in sections.get("conclusions", []):
        doc.add_paragraph(c)

    # 七、後續措施
    doc.add_heading("七、後續措施", level=2)
    for f in sections.get("follow_up", []):
        doc.add_paragraph(f)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def generate_report(cfg, output_path):
    """Generate the full incident report .docx document.

    Args:
        cfg: Configuration dict (see SKILL.md for schema)
        output_path: Path to write the .docx file

    Returns:
        The output path string
    """
    doc = Document()

    # Page setup: A4 with compact margins
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.45)
        section.bottom_margin = Cm(2.45)
        section.left_margin = Cm(1.99)
        section.right_margin = Cm(1.95)

    # Part 1: Government form
    build_form(doc, cfg)

    # Page break
    doc.add_page_break()

    # Part 2: Appendix
    build_appendix(doc, cfg)

    doc.save(output_path)
    print(f"Report generated: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description="Generate Taiwan cybersecurity incident report (個人資料侵害事故通報與紀錄表)")
    parser.add_argument("--output", "-o", required=True, help="Output .docx file path")
    parser.add_argument("--config", "-c", help="Config JSON file (see SKILL.md for schema)")
    args = parser.parse_args()

    cfg = dict(DEFAULTS)
    if args.config:
        with open(args.config) as f:
            user_cfg = json.load(f)
        cfg = deep_merge(cfg, user_cfg)

    generate_report(cfg, args.output)


if __name__ == "__main__":
    main()
